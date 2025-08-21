import os
import openai
from supabase import create_client
from dotenv import load_dotenv
import logging
import re
from notion_client import Client
from notion_util import extract_blocks_from_page 
import hashlib
import tiktoken
import json
import base64
import io
from PIL import Image as PILImage

# This script extracts content from Notion pages and attached files, embeds them using OpenAI embeddings,
# and stores the embeddings in a Supabase PostgreSQL vector database for semantic retrieval by Juno.

#imports for enhanced file processing
import fitz  # PyMuPDF for PDF parsing
try:
    import docx  # python-docx for DOCX files
except ImportError:
    docx = None
    logging.warning("python-docx not installed. DOCX support disabled.")

try:
    from PIL import Image
    import pytesseract
    import io
except ImportError:
    Image = None
    pytesseract = None
    logging.warning("PIL/pytesseract not installed. OCR support disabled.")

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load env variables
load_dotenv()
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Initialize clients
openai.api_key = OPENAI_API_KEY
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
notion = Client(auth=os.getenv("NOTION_TOKEN"))

def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from a PDF file, including image OCR if available."""
    extracted_text = []
    try:
        with fitz.open(file_path) as doc:
            for page_num in range(len(doc)):
                page = doc[page_num]

                #get regular text
                text = page.get_text()
                if text and text.strip(): 
                    extracted_text.append(f"Page {page_num + 1}:\n{text.strip()}\n")
                #get images and perform OCR if available
                if Image and pytesseract:
                    for img_index, img in enumerate(page.get_images(full=True)):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            image = Image.open(io.BytesIO(image_bytes))
                            image = image.convert("RGB")  # Convert to RGB for OCR compatibility
                            ocr_text = pytesseract.image_to_string(image)
                            if ocr_text.strip():
                                extracted_text.append(f"[Page {page_num + 1} - Image {img_index + 1}]\n{ocr_text.strip()}")
                        except Exception as e:
                            logger.warning(f"Error processing image on page {page_num + 1}: {e}")
        return "\n\n".join(extracted_text)
    except Exception as e:
        logger.error(f"Error extracting text from PDF file {file_path}: {e}")
        return "\n\n".join(extracted_text)

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a DOCX file."""
    if not docx:
        logger.warning("python-docx is not installed. Cannot process DOCX files.")
        return ""
    
    text_parts = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {file_path}: {e}")
    return "\n".join(text_parts)

def extract_text_from_image(file_path: str) -> str:
    """Extracts text from an image file using OCR."""
    if not Image or not pytesseract:
        logger.warning("PIL/pytesseract is not installed. Cannot process image files.")
        return ""
    
    try:
        image = Image.open(file_path)
        image.convert("RGB")  # Convert to RGB for better OCR compatibility
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from image file {file_path}: {e}")
        return ""

def enhanced_extract_text(file_path: str) -> str:
    """Enhanced text extractor that handles multiple file types."""
    file_type = file_path.lower()
    
    try:
        if file_type.endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
        elif file_type.endswith((".docx", ".doc")):
            text = extract_text_from_docx(file_path)
        elif file_type.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
            text = extract_text_from_image(file_path)
        elif file_type.endswith((".txt", ".md")):
            # Use existing logic for text files
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            #to debug repeating lines
            logger.info(f"[debug] {file_path} len={len(text)}")
            for i, line in enumerate(text.splitlines()[:15]):
                logger.info(f"[debug] line {i+1}: {line}")
            return text
            #
        else:
            # Try to read as text file (fallback)
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        if text is None:
            logger.warning(f"Text extraction returned None for {file_path}")
            return ""
        
        # Ensure we always return a string, even if empty
        return text if text is not None else ""
        
    except Exception as e:
        logger.error(f"Cannot process file type for {file_path}: {e}")
        return ""
        
def get_file_hash(file_path):
    """Generate a hash for the file to check if it has been processed"""
    try:
        with open(file_path, 'rb') as f:
            file_content = f.read()
            return hashlib.md5(file_content).hexdigest()
    except Exception as e:
        logger.error(f"Error generating hash for {file_path}: {e}")
        return None

def hash_chunk(text):
    """Generate a hash for a text chunk to check if it has been processed"""
    try:
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    except Exception as e:
        logger.error(f"Error generating hash for chunk: {e}")
        return None

#metadata helpers
def stable_doc_id(page_name: str | None) -> str | None:
    if not page_name:
        return None
    return hashlib.md5(page_name.encode("utf-8")).hexdigest()

def infer_source_kind(file_path: str | None, fallback: str = "document") -> str:
    if not file_path:
        return fallback
    ext = os.path.splitext(file_path.lower())[1]
    if ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif", ".webp", ".svg"):
        return "image"
    if ext in (".eml", ".msg"):
        return "email"
    return fallback

def ext_from_path(file_path: str | None) -> str | None:
    if not file_path:
        return None
    ext = os.path.splitext(file_path)[1].lower()
    return ext[1:] if ext.startswith(".") else ext

def quick_word_count(text: str) -> int:
    return len((text or "").split())

def parse_email_headers(text: str) -> dict | None:
    """
    If some of your 'documents' are email bodies saved as text,
    try to parse basic headers from the first ~20 lines.
    """
    lines = (text or "").splitlines()
    if not lines:
        return None
    headers = {}
    for line in lines[:20]:
        if ":" not in line:
            continue
        key, val = line.split(":", 1)
        key = key.strip().lower()
        val = val.strip()
        if key in ("from", "to", "subject", "date"):
            headers[key] = val
    return headers or None

#wont be used as much, going to be using is_chunk_already_embedded instead
# but keeping this for future reference
def is_file_already_embedded(file_path, page_name):
    """Check if a file has already been embedded by checking the database"""
    try:
        # Check by exact file path and page name
        result = supabase.table("juno_embeddings").select("id").eq("source_file", file_path).eq("page_name", page_name).limit(1).execute()
        
        if result.data:
            logger.info(f"File already embedded (exact match): {file_path}")
            return True
        
        # Check by file hash (for when files are re-downloaded with different paths)
        file_hash = get_file_hash(file_path)
        if file_hash:
            # Check if we have this hash in our database
            hash_result = supabase.table("juno_embeddings").select("id").eq("file_hash", file_hash).limit(1).execute()
            
            if hash_result.data:
                logger.info(f"File already embedded (hash match): {file_path}")
                return True
        
        return False
        
    except Exception as e:
        logger.error(f"Error checking if file is embedded: {e}")
        return False  # If we can't check, assume it's not embedded to be safe


def split_text(text, max_tokens=800, overlap_tokens=100):
    """Split text into overlapping chunks optimized for Q&A and robust for any text format
    Handles: paragraphs, massive single paragraphs, podcast transcripts, etc."""

    enc = tiktoken.get_encoding("cl100k_base") # Using the same encoding as OpenAI
    #clean input
    text = text.strip()
    if not text:
        return []
    
     # If the entire text is small enough, return as single chunk
    total_tokens = len(enc.encode(text))
    if total_tokens <= max_tokens:
        return [text]
    
    # For very large texts, use safe sliding window approach
    if total_tokens > 20000:  # If >20k tokens, use safe approach
        logger.info(f"Large text detected ({total_tokens} tokens), using sliding window")
        return _safe_sliding_split(text, max_tokens, overlap_tokens, enc)
    try:
        paragraphs = text.split("\n\n")  # Split by double newlines to get paragraphs
        chunks = []
        current_chunk = ""
        
        for i, para in enumerate(paragraphs):
            para = para.strip()
            if not para:
                continue
            
            # Safety check - don't let any single paragraph be too big
            para_tokens = len(enc.encode(para))
            if para_tokens > max_tokens * 2:  # If paragraph is huge, force split it
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                chunks.extend(_safe_sliding_split(para, max_tokens, overlap_tokens, enc))
                continue

            # Normal paragraph processing
            test_chunk = current_chunk + "\n\n" + para if current_chunk else para
            test_tokens = len(enc.encode(test_chunk))
            
            if test_tokens <= max_tokens:
                current_chunk = test_chunk
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = para
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
        
    except Exception as e:
        logger.warning(f"Paragraph splitting failed: {e}, falling back to sliding window")
        return _safe_sliding_split(text, max_tokens, overlap_tokens, enc)

def _safe_sliding_split(text, max_tokens, overlap_tokens, enc):
    """Safe sliding window split that can't get stuck"""
    try:
        tokens = enc.encode(text)
        chunks = []
        start = 0
        max_iterations = len(tokens) // (max_tokens - overlap_tokens) + 10  # Safety limit
        iterations = 0
        
        while start < len(tokens) and iterations < max_iterations:
            end = min(start + max_tokens, len(tokens))
            chunk_tokens = tokens[start:end]
            chunk_text = enc.decode(chunk_tokens).strip()
            
            if chunk_text:
                chunks.append(chunk_text)
            
            # Move forward with overlap
            start = end - overlap_tokens
            if start >= end or end >= len(tokens):
                break
                
            iterations += 1
        
        return chunks
        
    except Exception as e:
        logger.error(f"Even sliding window failed: {e}")
        # Last resort - just return first chunk
        try:
            first_tokens = enc.encode(text)[:max_tokens]
            return [enc.decode(first_tokens).strip()]
        except:
            return ["Error processing text"]

def embed_and_store(file_path, page_name):
    """Embed file content and store in Supabase, enhanced to handle various file types"""
    
    # Check if file exists
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return
    
    #gen file hash for duplicate file checking
    file_hash = get_file_hash(file_path)

    if is_image_file(file_path):
        logger.info(f"Processing image file: {file_path}")
        #check if image is already embedded
        result = supabase.table("juno_embeddings").select("id").eq("file_hash", file_hash).eq("page_name", page_name).limit(1).execute()
        if result.data:
            logger.info(f"Image already embedded: {file_path}")
            return
        # Store image in Supabase storage
        image_url = store_image(file_path, page_name)
        if not image_url:
            logger.error(f"Failed to store image: {file_path}")
            return
        description = get_image_description(file_path)
        ocr_text = ""
        try:
            if Image and pytesseract:
                ocr_text = extract_text_from_image(file_path)
        except:
            pass # If OCR fails, just skip it
        #combine descruption and OCR text
        searchable_text = f"Image: {description}"
        if ocr_text.strip():
            searchable_text += f"\n\nText in image: {ocr_text}"
        
        # Insert image metadata into Supabase
        try:
            response = openai.embeddings.create(
                input=searchable_text,
                model="text-embedding-3-small"
            ).data[0].embedding
            
            # metadata for images
            img_md = {
                "doc_id": stable_doc_id(page_name),
                "source_kind": "image",
                "title": page_name,
                "filename": os.path.basename(file_path) if file_path else None,
                "ext": ext_from_path(file_path),
                "word_count": quick_word_count(description or searchable_text),
            }
            # Store in database with image metadata
            supabase.table("juno_embeddings").insert({
                "page_name": page_name,
                "chunk": searchable_text,
                "embedding": response,
                "source_file": file_path,
                "chunk_index": 0,
                "file_hash": file_hash,
                "chunk_hash": hash_chunk(searchable_text),
                "file_type": "image",
                "image_url": image_url,  # Store the public URL
                "original_filename": os.path.basename(file_path),
                "description": description,  # Store the image description
                "metadata": img_md,
            }).execute()
            
            logger.info(f"Successfully embedded image: {file_path}")
            
        except Exception as e:
            logger.error(f"Error embedding image {file_path}: {e}")
    else:
        # Regular file processing (PDF, DOCX, TXT, etc.)
        try:
            text = enhanced_extract_text(file_path)
            #debug repeating lines
            logger.info(f"[debug-pre-chunk] {file_path} first 300 chars: {text[:300]!r}")
        
            if not text.strip():
                logger.warning(f"No text content found in {file_path}")
                return
            
        except Exception as e:
            logger.error(f"Error reading {file_path}: {e}")
            return
    
        # Split the full text into smaller chunks
        chunks = split_text(text)
        logger.info(f"Split {file_path} into {len(chunks)} chunks")
    
        # Embed each chunk and store in Supabase
        for i, chunk in enumerate(chunks):

            chunk_hash = hash_chunk(chunk) # Check if this chunk has already been embedded
            result = supabase.table("juno_embeddings").select("id").eq("chunk_hash", chunk_hash).eq("page_name", page_name).limit(1).execute()
            if result.data:
                logger.info(f"Chunk {i+1}/{len(chunks)} already embedded. Skipping.")
                continue
            if chunk.strip():
                try:
                    # Get embedding from OpenAI
                    response = openai.embeddings.create(
                        input=chunk,
                        model="text-embedding-3-small"
                    ).data[0].embedding
                
                    # Insert the chunk and its embedding into the Supabase table
                    supabase.table("juno_embeddings").insert({
                        "page_name": page_name,
                        "chunk": chunk,
                        "embedding": response,
                        "source_file": file_path,
                        "chunk_index": i,
                        "file_hash": file_hash,
                        "chunk_hash": chunk_hash,
                        "file_type": "document",
                        "original_filename": os.path.basename(file_path),
                    }).execute()
                
                    logger.info(f"Inserted chunk {i+1}/{len(chunks)} from {file_path}")
                
                except Exception as e:
                    logger.error(f"Error processing chunk {i} from {file_path}: {e}")

#should not be used in production, only for debugging because pages can have files added/removed
def is_page_already_processed(page_name):
    """Check if a Notion page has already been processed"""
    try:
        result = supabase.table("juno_embeddings").select("id").eq("page_name", page_name).limit(1).execute()
        return len(result.data) > 0
    except Exception as e:
        logger.error(f"Error checking if page is processed: {e}")
        return False

def get_page_title(page):
    """Extract title from Notion page object"""
    properties = page.get("properties", {})
    
    # Try different possible title property names
    for title_key in ["title", "Page", "Name"]:
        if title_key in properties:
            title_prop = properties[title_key]
            if title_prop.get("type") == "title" and title_prop.get("title"):
                return title_prop["title"][0]["plain_text"]
    
    # Fallback: try to get title from the page object itself
    if "title" in page:
        title_list = page["title"]
        if title_list and len(title_list) > 0:
            return title_list[0]["plain_text"]
    
    # Last resort: use page ID
    return f"Page_{page['id']}"

def batch_embed_all_notion_files(force_reprocess=False):#force_reprocess=False incase of overwriting existing data
    """Search all Notion pages for files, embed them, and store in Supabase, enhanced with better file support"""
    logger.info("Connecting to Notion...")
    
    try:
        # Get all pages from Notion
        search_results = notion.search()["results"]
        logger.info(f"Found {len(search_results)} pages from Notion search")
        
        total_files_processed = 0
        
        for page in search_results:
            page_id = page["id"]
            page_title = get_page_title(page)
            
            logger.info(f"Processing page: '{page_title}' (ID: {page_id})")
            
            try:
                # Extract content directly using page ID
                content = extract_blocks_from_page(page_id)
                
                if not content:
                    logger.warning(f"No content found for page: {page_title}")
                    continue
                
                # Look for downloaded files in the content
                file_paths = re.findall(r"\[FILE_DOWNLOADED\]:(.+)", content)
                
                if file_paths:
                    logger.info(f"Found {len(file_paths)} files in page '{page_title}'")
                    
                    for file_path in file_paths:
                        file_path = file_path.strip()
                        logger.info(f"Processing file: {file_path}")
                        
                        embed_and_store(file_path, page_title)
                        total_files_processed += 1
                else:
                    logger.info(f"No files found in page '{page_title}'")
                    
                    # Also embed the page content itself if it has text
                    text_content = re.sub(r"\[FILE_DOWNLOADED\]:[^\n]*", "", content).strip()
                    if text_content:
                        # Create a temporary text file for the page content
                        temp_file = f"temp_page_{page_id}.txt"
                        try:
                            with open(temp_file, "w", encoding="utf-8") as f:
                                f.write(text_content)
                            logger.info(f"Created temporary file for page content: {temp_file}")
                            embed_and_store(temp_file, page_title)
                            total_files_processed += 1
                            os.remove(temp_file)  # Clean up
                        except Exception as e:
                            logger.error(f"Error processing page content for {page_title}: {e}")
                
            except Exception as e:
                logger.error(f"Error processing page '{page_title}': {e}")
                continue
        
        logger.info(f"Batch processing complete. Total files processed: {total_files_processed}")
        
    except Exception as e:
        logger.error(f"Error in batch processing: {e}")

def test_single_page(page_name):
    """Test function to debug a single page"""
    logger.info(f"Testing single page: {page_name}")
    
    # Search for the specific page
    search_results = notion.search(query=page_name)["results"]
    
    if not search_results:
        logger.error(f"Page '{page_name}' not found")
        return
    
    page = search_results[0]
    page_id = page["id"]
    page_title = get_page_title(page)
    
    logger.info(f"Found page: '{page_title}' (ID: {page_id})")
    
    # Extract content
    content = extract_blocks_from_page(page_id)
    logger.info(f"Page content preview: {content[:200]}...")
    
    # Look for files
    file_paths = re.findall(r"\[FILE_DOWNLOADED\]:(.+)", content)
    logger.info(f"Found {len(file_paths)} files: {file_paths}")
    
    return content

def is_image_file(file_path: str) -> bool:
    """Check if a file is an image based on its extension"""
    image_extensions = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif", ".webp", ".svg"}
    return any(file_path.lower().endswith(ext) for ext in image_extensions)
def store_image(file_path: str, page_name: str) -> str:
    """Store image in Supabase storage and return public URL"""
    try:
        # Generate unique filename
        file_hash = get_file_hash(file_path)
        file_ext = os.path.splitext(file_path)[1]
        storage_filename = f"{file_hash}{file_ext}"
        
        # Read image file
        with open(file_path, 'rb') as f:
            file_bytes = f.read()
        
        # Upload to Supabase storage
        supabase.storage.from_("juno_images").upload(
            path=storage_filename,
            file=file_bytes,
            file_options={"x-upsert": "true"}
        )
        
        # Get public URL
        public_url = f"{SUPABASE_URL}/storage/v1/object/public/juno_images/{storage_filename}"
        
        logger.info(f"Stored image in Supabase: {public_url}")
        return public_url
        
    except Exception as e:
        logger.error(f"Error storing image in Supabase: {e}")
        return None
def get_image_description(file_path: str) -> str:
    """Generate description of image using OpenAI Vision API"""
    try:
        # Convert image to base64
        with open(file_path, 'rb') as f:
            image_bytes = f.read()
        
        # Resize if too large (OpenAI has size limits)
        image = PILImage.open(io.BytesIO(image_bytes))
        if image.width > 2000 or image.height > 2000:
            image.thumbnail((2000, 2000), PILImage.Resampling.LANCZOS)
            buffer = io.BytesIO()
            image.save(buffer, format='JPEG', quality=85)
            image_bytes = buffer.getvalue()
        
        base64_image = base64.b64encode(image_bytes).decode('utf-8')
        
        # Get description from OpenAI
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Describe this image in detail, including colors, objects, text, style, and any other relevant visual elements. This will be used for semantic search. Keep it to one sentence if possible but include all important details. Is it a logo, picture, or something else?"
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=500
        )
        
        description = response.choices[0].message.content
        logger.info(f"Generated image description: {description[:100]}...")
        return description
        
    except Exception as e:
        logger.error(f"Error generating image description: {e}")
        return f"Image file: {os.path.basename(file_path)}"

def search_for_image(query: str, limit: int = 5):
    """Search for images that match a query for Juno to use. 1. direct text search in file description, 
    2. ocr text in images, 3. image similarity search using OpenAI embeddings"""
    try:
        all_results = []
        seen_urls = set()  # Prevent duplicates
        
        # Method 1: Direct text search in description (FASTEST - SQL LIKE query)
        description_results = supabase.table("juno_embeddings").select(
            "page_name, description, image_url, original_filename, chunk"
        ).eq('file_type', 'image').ilike('description', f'%{query}%').limit(limit).execute()
        
        for match in description_results.data:
            if match['image_url'] not in seen_urls:
                all_results.append({
                    'page_name': match['page_name'],
                    'description': match['description'],
                    'image_url': match['image_url'],
                    'filename': match['original_filename'],
                    'similarity': 0.95,  # High confidence for exact text matches
                    'match_type': 'description_match',
                    'match_text': query
                })
                seen_urls.add(match['image_url'])
        
        # Method 2: Search in OCR text (for text found within images)
        # Only search if we haven't found enough results yet
        if len(all_results) < limit:
            remaining_limit = limit - len(all_results)
            ocr_results = supabase.table("juno_embeddings").select(
                "page_name, description, image_url, original_filename, chunk"
            ).eq('file_type', 'image').ilike('chunk', f'%Text in image: %{query}%').limit(remaining_limit).execute()
            
            for match in ocr_results.data:
                if match['image_url'] not in seen_urls:
                    all_results.append({
                        'page_name': match['page_name'],
                        'description': match['description'],
                        'image_url': match['image_url'],
                        'filename': match['original_filename'],
                        'similarity': 0.90,  # High confidence for OCR matches
                        'match_type': 'ocr_match',
                        'match_text': query
                    })
                    seen_urls.add(match['image_url'])
        
        # Method 3: Semantic search using embeddings (SLOWER but finds conceptual matches)
        # Only if we still need more results
        if len(all_results) < limit:
            remaining_limit = limit - len(all_results)
            
            # Get embedding for the search query
            response = openai.embeddings.create(
                input=query,
                model="text-embedding-3-small"
            ).data[0].embedding
            
            # Semantic search
            semantic_results = supabase.rpc('match_documents', {
                'query_embedding': response,
                'match_threshold': 0.6,  # Lower threshold for broader matches
                'match_count': remaining_limit
            }).eq('file_type', 'image').execute()
            
            for match in semantic_results.data:
                if match['image_url'] not in seen_urls:
                    all_results.append({
                        'page_name': match['page_name'],
                        'description': match.get('description', match['chunk']),
                        'image_url': match['image_url'],
                        'filename': match['original_filename'],
                        'similarity': match.get('similarity', 0.7),
                        'match_type': 'semantic_match',
                        'match_text': None
                    })
                    seen_urls.add(match['image_url'])
        
        # Sort by similarity score (highest first)
        all_results.sort(key=lambda x: x['similarity'], reverse=True)
        
        return all_results[:limit]
        
    except Exception as e:
        logger.error(f"Error searching images: {e}")
        return []

if __name__ == "__main__":
    # Uncomment the next line to test a specific page first
    # test_single_page("Competition Handbook & Rules")
    logger.info("Starting enhanced embedder with PDF and multipart file support...")
    
    # Check for required dependencies
    missing_deps = []
    if not fitz:
        missing_deps.append("PyMuPDF (pip install PyMuPDF)")
    if not docx:
        missing_deps.append("python-docx (pip install python-docx)")
    if not Image or not pytesseract:
        missing_deps.append("PIL and pytesseract (pip install Pillow pytesseract)")
    
    if missing_deps:
        logger.warning(f"Missing optional dependencies: {', '.join(missing_deps)}")
        logger.warning("Some file types may not be processed correctly.")
    
    batch_embed_all_notion_files(force_reprocess=False)  # Set to True to reprocess all pages
    # Note: This will process all Notion pages and embed their content/files into Supabase